"""DOCX extractor using python-docx + PaddleOCR-VL v1.5.

Extracts text from paragraphs and tables via python-docx.
Embedded images are exported and sent to PaddleOCR-VL for OCR.
"""

import logging
import tempfile
from pathlib import Path

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT

from ingest.ocr_engine import get_engine

logger = logging.getLogger(__name__)


def _extract_paragraphs(doc: Document) -> list[str]:
    """Extract text from all paragraphs."""
    return [p.text for p in doc.paragraphs if p.text.strip()]


def _extract_tables(doc: Document) -> list[str]:
    """Extract text from all tables as markdown tables."""
    tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        if not rows:
            continue
        # Markdown table format
        header = "| " + " | ".join(rows[0]) + " |"
        sep = "| " + " | ".join("---" for _ in rows[0]) + " |"
        body = "\n".join("| " + " | ".join(r) + " |" for r in rows[1:])
        tables.append(f"{header}\n{sep}\n{body}")
    return tables


def _extract_images(doc: Document, docx_path: Path) -> list[str]:
    """Extract embedded images and OCR them via PaddleOCR-VL v1.5."""
    results = []
    engine = get_engine()

    for i, rel in enumerate(doc.part.rels.values()):
        if "image" not in rel.reltype:
            continue
        try:
            img_data = rel.target_part.blob
            suffix = Path(rel.target_ref).suffix or ".png"
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as f:
                f.write(img_data)
                tmp_path = Path(f.name)

            logger.info("OCR-ing embedded image #%d from %s", i + 1, docx_path.name)
            text = engine.ocr_file(tmp_path)
            if text.strip():
                results.append(text)
        except Exception:
            logger.exception("Failed to OCR embedded image #%d", i + 1)
        finally:
            try:
                tmp_path.unlink()
            except Exception:
                pass

    return results


def extract(docx_path: str | Path) -> str:
    """Extract text from DOCX via python-docx + PaddleOCR-VL v1.5.

    Returns raw text: paragraphs + tables + OCR'd images.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"DOCX not found: {docx_path}")

    doc = Document(str(docx_path))

    parts = []
    paragraphs = _extract_paragraphs(doc)
    if paragraphs:
        parts.append("\n".join(paragraphs))

    tables = _extract_tables(doc)
    if tables:
        parts.append("\n\n".join(tables))

    images = _extract_images(doc, docx_path)
    if images:
        parts.append("\n\n".join(images))

    return "\n\n".join(parts)
